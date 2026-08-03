# -*- coding: utf-8 -*-
"""Apply docs/report markdown updates into the PillSafe FYP Word document."""
from __future__ import annotations

from copy import deepcopy

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph

DOC_PATH = r"C:\Users\Boison\Desktop\PillSafe Final Year Project.docx"
PENDING = "Pending measurement"


def set_runs_text(paragraph: Paragraph, text: str) -> None:
    """Replace paragraph text while keeping the paragraph style."""
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def delete_paragraph(paragraph: Paragraph) -> None:
    el = paragraph._element
    parent = el.getparent()
    if parent is not None:
        parent.remove(el)


def add_para(doc: Document, text: str, style: str = "Normal") -> Paragraph:
    p = doc.add_paragraph(text)
    try:
        p.style = style
    except KeyError:
        pass
    return p


def add_heading_para(doc: Document, text: str, level: int) -> Paragraph:
    style = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3"}.get(level, "Normal")
    return add_para(doc, text, style)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            table.rows[r + 1].cells[c].text = val
    doc.add_paragraph("")


def global_replace_in_paragraphs(doc: Document, replacements: list[tuple[str, str]]) -> int:
    count = 0
    for p in doc.paragraphs:
        text = p.text
        if not text:
            continue
        new = text
        for old, rep in replacements:
            if old in new:
                new = new.replace(old, rep)
        if new != text:
            set_runs_text(p, new)
            count += 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    text = p.text
                    if not text:
                        continue
                    new = text
                    for old, rep in replacements:
                        if old in new:
                            new = new.replace(old, rep)
                    if new != text:
                        set_runs_text(p, new)
                        count += 1
    return count


ABSTRACT = (
    "Medication non-adherence remains a major cause of treatment failure, especially among "
    "elderly and chronically ill patients who manage multiple daily doses. Existing electronic "
    "pill dispensers often lack strong identity verification, rely on weak alarms, or do not "
    "notify caregivers when a dose is missed. This project presents PillSafe, a biometric "
    "medication-dispensing system that combines a Raspberry Pi 5 hub with a React Native mobile "
    "application.\n\n"
    "The hub stores prescribed schedules in SQLite, reminds the patient when a dose becomes due "
    "(buzzer on the device and an in-app alert), and releases medication only after successful "
    "biometric verification. Authentication uses facial recognition (MobileFaceNet TensorFlow "
    "Lite with OpenCV Haar detection) and an optional voice-enrolment / verification path via a "
    "Voice HAT microphone. Six patient compartments are driven by MG996R continuous-rotation "
    "servos; each compartment holds nine angular slots at 40° steps. After a successful verify, "
    "the assigned servo advances one slot and gravity delivers the dose; optional IR sensors can "
    "confirm drop and pickup. Missed doses after a configurable grace period, repeated "
    "verification failures, and related faults generate SMS alerts to a caregiver through a "
    "SIM800C module on the Pi UART. The phone app manages device connection, schedules "
    "(including AM/PM dosing and edit/delete), live camera preview for enrolment and verify, "
    "adherence monitoring, and the Verify Now handshake with the hub.\n\n"
    "The system is designed for home and small-clinic use where reliable, identity-gated "
    "dispensing and caregiver visibility matter more than cloud-centric pharmacy automation. "
    "Implementation, construction, and evaluation are reported in Chapters 4–6.\n\n"
    "Keywords: medication adherence, facial recognition, voice biometrics, Raspberry Pi 5, "
    "React Native, GSM SMS, smart pill dispenser."
)

# Specific paragraph rewrites keyed by unique old substring prefixes (matched once).
SPECIFIC_REWRITES: list[tuple[str, str]] = [
    (
        "Hardware Layer contains all physical elements within the PillSafe device that deal with sensing, dispensing, and communication operations. It includes Raspberry Pi 4B",
        "Hardware Layer contains all physical elements within the PillSafe device that deal with sensing, dispensing, communication, and optional voice capture. It includes a Raspberry Pi 5, Pi Camera Module, six MG996R continuous-rotation servos (one per patient compartment), FC-51 infrared sensors, an active buzzer, a DS3231 RTC, and a SIM800C GSM module on the Pi UART, with an optional Voice HAT / I2S microphone path.",
    ),
    (
        "Application & Monitoring Layer consists of the mobile app built with Flutter",
        "Application & Monitoring Layer consists of the React Native mobile application (with an Expo Go variant for camera-less UI preview during development) that communicates with the Raspberry Pi hub over a local Wi-Fi hotspot or LAN via REST, including device connection with an API token, schedule management, reminders, verify handshake, and adherence monitoring.",
    ),
    (
        "In case there is a match between a medication time and the schedule data in the database, the system triggers an alarm signal using the buzzer component and initiates the facial ve",
        "When a medication time matches the schedule, the hub runs a dose-due buzzer reminder pattern and the mobile app can surface a “Time Is Up” alert via its reminder poller. The buzzer is stopped before dispense authentication so it does not interfere with camera or microphone capture. Dispensing proceeds only after the app issues Verify Now (when require_verify_request is enabled) and biometric verification succeeds.",
    ),
    (
        "Upon successful authentication, the Raspberry Pi component rotates the SG90 servo motor",
        "Upon successful authentication (face, with optional voice when enabled), the Raspberry Pi advances the assigned compartment’s MG996R continuous-rotation servo by one calibrated 40° slot step so that medication is released by gravity into the collection path. Optional IR sensors may confirm drop and pickup.",
    ),
    (
        "Records of dispensing processes and adherence details are stored locally in the SQLite database. In case verification fails or medication is not retrieved within the permissible ti",
        "Records of dispensing processes and adherence details are stored locally in the SQLite database. If verification fails repeatedly or the dose is not completed within the configurable grace period, the event is logged as MISSED (or related failure) and SMS alerts are sent to the caregiver through the SIM800C UART interface.",
    ),
    (
        "The facial recognition component was intended for authentication of the user prior to dispensing the medicine. The component includes the Pi Camera Module v2, pre-processing via Op",
        "The primary authentication path is facial recognition prior to dispensing. It uses the Pi Camera Module, OpenCV Haar detection, and MobileFaceNet TensorFlow Lite embeddings on the Raspberry Pi 5. Live MJPEG stream and JPEG snapshot APIs support phone-side preview during enrolment and verify. An optional secondary voice enrolment/verification path is provided via Voice HAT / I2S when enabled in configuration.",
    ),
    (
        "The dispensing unit consists of a six-compartment circular carousel mechanism actuated by an SG90 servo motor.",
        "The dispensing unit comprises six patient compartments, each with a rotating cylinder of nine angular dose slots at 40° (9 × 40° = 360°). Each compartment is actuated by its own MG996R continuous-rotation servo. Delivery is gravity-based after the slot advances; there is no motorised gate in the baseline design. CAD sizing targets a compartment envelope on the order of Ø200 mm × 21.67 mm height, with approximately 72 teeth on the large layer gear and a matched 16-tooth pinion.",
    ),
    (
        "Medication for each particular user is kept in a distinct compartment. Following successful user authentication, the Raspberry Pi sends a PWM command signal to the servo motor to r",
        "Medication for each patient is kept in a distinct compartment. Following successful authentication, the Raspberry Pi drives the corresponding servo with a timed PWM duty around a calibrated neutral pulse (continuous mode), advancing software slot index by one rather than seeking an absolute 0–360° positional angle.",
    ),
    (
        "The required angular displacement of each compartment can be computed using the formula below:",
        "The angular step per dose slot within a compartment is fixed by geometry:",
    ),
    (
        "= angular separation between compartments",
        "= angular separation between slots in one compartment",
    ),
    (
        "= total number of compartments",
        "= total number of slots per compartment (nine)",
    ),
    (
        "For six compartments:",
        "For nine slots:",
    ),
    (
        "The Servo rotates in steps of 60 degrees to position the appropriate compartment directly under the dispensing chute.",
        "Each compartment servo rotates in calibrated steps of 40 degrees to advance to the next dose slot. Continuous-rotation timing is tuned via degrees_per_second, neutral_duty, and run_duty_offset in config.yaml so that one command corresponds to approximately one slot.",
    ),
    (
        "Pulse Width Modulation duty cycle was regulated using the Raspberry Pi GPIO hardware PWM interface. The dispensing time is calculated using the following equation:",
        "Pulse-width modulation for the six servos is generated on BCM pins 12, 13, 22, 17, 26 and 27 (GPIO 16 is avoided because of Voice HAT conflict). A PCA9685 expander is not required for the implemented pin map. Step duration is derived from the configured angle per slot and degrees_per_second calibration rather than a single shared 60° carousel formula.",
    ),
    (
        "The GSM communication module was designed using a SIM800L GSM module that was interfaced to the Raspberry Pi computer via USB-to-serial communication.",
        "The GSM communication module uses a SIM800C interfaced to the Raspberry Pi UART (/dev/serial0 at 9600 baud): SIM800C TX to GPIO15 (RX), RX to GPIO14 (TX), with a separate 3.7–4.2 V LiPo supply and common ground with the Pi.",
    ),
    (
        "The Raspberry Pi computer uses serial AT commands to communicate with the SIM800L GSM module.",
        "The Raspberry Pi uses serial AT commands to communicate with the SIM800C GSM module.",
    ),
    (
        "Figure 8: The SIM800L GPRS GSM Module connected to the Raspberry Pi",
        "Figure 8: The SIM800C GSM module connected to the Raspberry Pi UART",
    ),
    (
        "The servo motor was interfaced using GPIO18 pin for PWM signal generation to control the servo motor. On the other hand, the two FC-51 sensors were interfaced using GPIO23 and GPIO24 pins.",
        "The six MG996R servos are interfaced on BCM pins 12, 13, 22, 17, 26 and 27 for PWM. The two FC-51 infrared sensors use GPIO23 (drop) and GPIO24 (pickup). The active buzzer uses GPIO25 and sounds only for dose-due reminders (stopped before dispense verify).",
    ),
    (
        "The RTC Module DS3231 connects to Raspberry Pi via the I2C Communication Interface by GPIO2 (SDA) and GPIO3 (SCL). SIM800L GSM connects via the Serial USB connection interface.",
        "The RTC Module DS3231 connects to the Raspberry Pi via I2C on GPIO2 (SDA) and GPIO3 (SCL). The SIM800C GSM module connects via the Pi UART (/dev/serial0), not USB-serial.",
    ),
    (
        "Figure 10: Raspberry pi 4B pins labelled",
        "Figure 10: Raspberry Pi 5 GPIO pins used by PillSafe",
    ),
    (
        "The Raspberry Pi 4B is responsible for serving as the embedded processor to control all the hardware components of PillSafe system.",
        "The Raspberry Pi 5 is responsible for serving as the embedded processor to control all the hardware components of the PillSafe system. GPIO access uses a compatibility layer (lgpio preferred, with RPi.GPIO fallback and a simulation backend for dry-run).",
    ),
    (
        "Flask API is hosted by the Raspberry Pi 4B to enable the embedded system to communicate with the mobile application that runs on the hotspot of Raspberry Pi 4B.",
        "A Flask REST API is hosted by the Raspberry Pi 5 to enable communication with the React Native mobile application over the Pi hotspot or LAN, including health, camera stream/snapshot, user and biometric enrolment, schedules CRUD, dispense request/verify, adherence, notifications, and dose-due test endpoints.",
    ),
    (
        "Powering the Raspberry Pi 4B occurs through an input regulated voltage of 5V. The SG90 servo motor also requires a 5V power source due to its rotational load through the PCA9685 16",
        "The Raspberry Pi 5 is powered by the official 27 W USB-C PD supply. The six MG996R servos use a separate external 5 V (≥5–6 A) rail with common ground to the Pi; they must not be powered from the Pi 5 V pins. The SIM800C uses a separate LiPo. A bulk capacitor on the servo rail absorbs stall spikes. Direct GPIO PWM is used; a PCA9685 board is not required for the implemented design.",
    ),
    (
        "Figure 11: The PCA9685 16-Channel PWM Servo Driver Board interfacing with Raspberry Pi",
        "Figure 11: Power domains for Pi, servo rail, and SIM800C (common ground; no PCA9685 required)",
    ),
    (
        "The design of the dispensing system was developed as a rotating six-pill compartments carousel manufactured using lightweight plastic through 3D printing technique.",
        "The dispensing structure is a six-layer / six-patient stack manufactured using lightweight plastic through 3D printing. Each layer provides nine dose slots. Sized Blender CAD defines compartment geometry and a drive gear (~72T) meshing with a 16T pinion.",
    ),
    (
        "The servo motor drives the carousel so that the selected compartment faces the dispensing hole; thereafter, the pills fall into the dispensing hole through which the medication dro",
        "The selected compartment’s servo advances one slot; thereafter pills fall by gravity into the delivery path toward the collection point, where optional IR sensing can confirm drop and pickup.",
    ),
    (
        "PillSafe mobile application is built on the Flutter development platform using Dart programming language.",
        "The PillSafe mobile application is built with React Native (TypeScript/JavaScript). An Expo Go project supports camera-less UI preview during development. It interacts with the Raspberry Pi over a local Wi-Fi hotspot or LAN via REST API endpoints secured with a shared API token.",
    ),
    (
        "The SMS alert system was implemented using the SIM800L GSM module and Python serial communication libraries.",
        "The SMS alert system was implemented using the SIM800C GSM module on /dev/serial0 and Python serial communication libraries.",
    ),
    (
        "The hardware architecture, which revolved around the Raspberry Pi 4B, was explained considering the computational requirements of the face recognition algorithm used at the edge.",
        "The hardware architecture, which revolves around the Raspberry Pi 5, was explained considering edge face recognition, optional voice, continuous MG996R slot indexing (six compartments × nine slots at 40°), SIM800C UART alerts, dose-due buzzer policy, and the React Native application. The next chapters present implementation/construction, testing results, and conclusions.",
    ),
    (
        "The operation of the embedded system starts with checking the medication RTC schedule. Whenever the time coincides with the scheduled time for medications, a signal is given throug",
        "The embedded system checks the RTC against medication schedules. When a dose becomes due, the dose-due buzzer runs and the app may alert the user. After Verify Now and successful face (± voice) verification, the dispenser advances the correct compartment one slot; IR sensors may confirm discharge and pickup. Adherence is updated in SQLite.",
    ),
    (
        "In the case of successful verification, the dispensing module turns the carousel, and infrared sensors detect whether the medicine has been dispensed and picked up.",
        "In the case of successful verification, the dispensing module advances the assigned continuous-rotation servo by one 40° slot, and infrared sensors may detect whether the medicine has been dispensed and picked up.",
    ),
    (
        "This project is separated into five chapters.",
        "This project is separated into six chapters.",
    ),
]


def apply_specific_rewrites(doc: Document) -> int:
    applied = 0
    for p in doc.paragraphs:
        text = p.text or ""
        if not text:
            continue
        for old_prefix, new_text in SPECIFIC_REWRITES:
            if text.startswith(old_prefix[:80]) or old_prefix[:60] in text[:120] or text.startswith(old_prefix):
                # Prefer exact startswith on a substantial prefix
                if text.startswith(old_prefix) or text.startswith(old_prefix[:100]):
                    set_runs_text(p, new_text)
                    applied += 1
                    break
                # Fallback: if unique short marker is present and paragraph is the known one
                if old_prefix[:90] in text:
                    set_runs_text(p, new_text)
                    applied += 1
                    break
    return applied


def update_hardware_table(doc: Document) -> None:
    if len(doc.tables) < 2:
        return
    table = doc.tables[1]
    rows = [
        ("Component", "Description"),
        (
            "Raspberry Pi 5",
            "High-performance single-board computer (64-bit OS) with gpio_compat (lgpio / RPi.GPIO / simulation) running face/voice pipelines, schedule control, dispenser drivers, Flask API, and SQLite.",
        ),
        (
            "Raspberry Pi Camera Module",
            "CSI camera for enrolment and verification; MJPEG stream and JPEG snapshot exposed to the mobile app.",
        ),
        (
            "MG996R continuous-rotation servos (×6)",
            "One servo per patient compartment; timed 40° slot steps with calibrated neutral duty and degrees_per_second.",
        ),
        (
            "FC-51 Infrared Sensors (×2)",
            "Infrared obstacle sensors used to detect pill drop and pickup (may be optional during bring-up).",
        ),
        (
            "DS3231 RTC Module",
            "Real-time clock module used for scheduling accuracy.",
        ),
        (
            "SIM800C GSM Module",
            "GSM module on Pi UART (/dev/serial0) for caregiver SMS on missed doses and verification failures; separate LiPo power.",
        ),
        (
            "Active Buzzer Module",
            "Dose-due reminder only; stopped before dispense authentication.",
        ),
        (
            "3D-printed six-compartment / nine-slot structure",
            "Gravity-driven stack (~Ø200 × 21.67 mm reference layer) with ~72T gear and matched 16T pinion.",
        ),
        (
            "Power supplies",
            "Official Pi 27 W USB-C PD; external 5 V ≥5–6 A for servos; LiPo for SIM800C; common ground bus.",
        ),
    ]
    # Resize if needed
    while len(table.rows) < len(rows):
        table.add_row()
    for i, (a, b) in enumerate(rows):
        table.rows[i].cells[0].text = a
        table.rows[i].cells[1].text = b


def update_comparison_table(doc: Document) -> None:
    if not doc.tables:
        return
    t = doc.tables[0]
    # Verification row
    for row in t.rows:
        if row.cells[0].text.strip().lower().startswith("verification"):
            row.cells[4].text = "Facial + optional voice"
        if row.cells[0].text.strip().lower().startswith("reminder"):
            row.cells[4].text = "Buzzer, SMS, and App"


def clear_from_chapter4(doc: Document) -> None:
    start = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith("Chapter 4"):
            start = i
            break
    if start is None:
        return
    # Delete from end to start so indices remain valid in the XML list we snapshot
    targets = list(doc.paragraphs[start:])
    for p in targets:
        delete_paragraph(p)


def append_chapters_4_to_6(doc: Document) -> None:
    # --- Chapter 4 ---
    add_heading_para(doc, "Chapter 4 - Implementation and Construction", 1)
    add_heading_para(doc, "4.1 Introduction", 2)
    add_para(
        doc,
        "This chapter describes how the PillSafe design was realised in hardware, embedded software, "
        "and the mobile client. It covers the bill of materials and wiring practice, the organisation "
        "of hub and app modules, biometric enrolment procedures, and the mechanical CAD decisions that "
        "support the six-compartment, nine-slot dispenser.",
    )

    add_heading_para(doc, "4.2 Bill of materials and wiring", 2)
    add_heading_para(doc, "4.2.1 Bill of materials (summary)", 3)
    add_para(
        doc,
        "The authoritative parts list is maintained in the project BOM (hardware/docs/bom.csv). "
        "Major items are summarised below.",
    )
    add_table(
        doc,
        ["Category", "Component", "Qty", "Role"],
        [
            ["Controller", "Raspberry Pi 5", "1", "Hub computer"],
            ["Power (Pi)", "Official 27 W USB-C PD", "1", "Pi only"],
            ["Storage", "microSD 32 GB+", "1", "Raspberry Pi OS"],
            ["Vision", "Pi Camera Module + CSI ribbon", "1", "Enrol / verify"],
            ["Dispensing", "MG996R continuous / 360° servo", "6", "One per compartment"],
            ["Dispensing", "External 5 V ≥5–6 A PSU", "1", "Servo VCC only"],
            ["Sensors", "FC-51 IR modules", "2", "Drop + pickup"],
            ["Feedback", "Active 5 V buzzer", "1", "Dose-due reminder"],
            ["Time", "DS3231 + CR2032", "1", "RTC"],
            ["GSM", "SIM800C + LiPo + SIM + antenna", "1 set", "Caregiver SMS"],
            ["Voice (optional)", "Voice HAT / I2S mic", "0–1", "Voice enrol/verify"],
            ["Passives", "Rail capacitors / optional series R", "as BOM", "Electrical safety"],
        ],
    )

    add_heading_para(doc, "4.2.2 Power architecture", 3)
    add_para(
        doc,
        "Three isolated positive rails share a single common ground: (1) Pi — official USB-C PD only; "
        "(2) servos — external 5 V high-current supply (never the Pi 5 V pins for six MG996R); "
        "(3) SIM800C — 3.7–4.2 V LiPo. A bulk electrolytic capacitor (e.g. 1000 µF / 16 V) across the "
        "servo rail absorbs stall spikes. Power-up order: wiring complete → GND verified → servo PSU → "
        "LiPo → Pi last.",
    )
    add_para(doc, "[Insert Figure 4.2 — power domains diagram]", "Caption")

    add_heading_para(doc, "4.2.3 Signal wiring summary", 3)
    add_para(
        doc,
        "Detailed pin-by-pin instructions are documented in the project wiring guide. Summary:",
    )
    add_table(
        doc,
        ["Function", "BCM / interface", "Notes"],
        [
            ["Servo PWM", "12, 13, 22, 17, 26, 27", "One line per compartment; GPIO 16 reserved / Voice HAT"],
            ["IR drop / pickup", "23 / 24", "3.3 V logic"],
            ["Buzzer", "25", "VCC from Pi 5 V if module is 5 V active"],
            ["RTC", "I2C (SDA/SCL)", "DS3231 @ 0x68"],
            ["GSM", "UART /dev/serial0", "SIM800C TX→15, RX→14; divider if TX > 3.3 V"],
            ["Camera", "CSI", "Stream + snapshot via API"],
            ["Voice", "Voice HAT / I2S", "Keep voice.enabled false until fitted"],
        ],
    )
    add_para(doc, "[Insert Figure 4.3 — annotated 40-pin header map]", "Caption")

    add_heading_para(doc, "4.3 Hub software implementation", 2)
    add_heading_para(doc, "4.3.1 Module map", 3)
    add_table(
        doc,
        ["Module / area", "Responsibility"],
        [
            ["main.py", "Process entry, wiring of controllers, API start"],
            ["config.yaml", "Single source of tuneable thresholds and pins"],
            ["gpio_compat.py", "Pi 5–friendly GPIO backends"],
            ["Face pipeline", "Haar + MobileFaceNet TFLite match"],
            ["Voice modules", "Challenge, enrol, verify (when enabled)"],
            ["Camera helpers", "Shared camera for MJPEG / JPEG endpoints"],
            ["dispenser.py", "Continuous servo slot steps, IR wait policy"],
            ["Schedule controller", "Due / grace / MISSED / REMINDER coordination"],
            ["gsm.py", "SIM800C AT commands over UART"],
            ["Buzzer driver", "Dose-due pattern; stopped before dispense"],
            ["api/routes.py", "Flask REST surface for the phone"],
            ["SQLite DB layer", "Users, schedules, adherence, inventory, notifications"],
        ],
    )
    add_para(doc, "[Insert Figure 4.4 — module block diagram]", "Caption")

    add_heading_para(doc, "4.3.2 Dispenser control", 3)
    add_para(
        doc,
        "With continuous servo mode, a dispense command does not seek an absolute PWM angle. It selects "
        "the compartment pin, runs the servo away from neutral_duty for a duration derived from "
        "angle_per_slot and degrees_per_second, returns to neutral, advances the software slot index, "
        "and optionally waits for IR confirmation. Calibration adjusts degrees_per_second until one "
        "step visually matches approximately 40° on the printed cylinder.",
    )

    add_heading_para(doc, "4.3.3 Scheduling and reminder policy", 3)
    add_para(
        doc,
        "Schedules and adherence events persist in SQLite. The hub polls on a short interval (default "
        "15 s). On dose-due, the buzzer pattern runs (~50 s typical). When require_verify_request is "
        "true, the hub waits for POST /dispense/request from the app before authentication. After "
        "success, dispense runs and adherence is updated; after grace expiry, MISSED and SMS paths fire.",
    )

    add_heading_para(doc, "4.3.4 REST API (implemented)", 3)
    add_para(
        doc,
        "The Flask service binds 0.0.0.0:5000 by default with a shared token. Implemented groups include "
        "health, camera snapshot/stream, user CRUD and enrolment, voice challenge/enrol, dispense "
        "request/verify, schedule CRUD, adherence, inventory, notifications, and dose-due test.",
    )

    add_heading_para(doc, "4.4 Mobile application implementation", 2)
    add_heading_para(doc, "4.4.1 Stack", 3)
    add_para(
        doc,
        "The production-facing client is React Native. A parallel Expo Go project allows rapid UI "
        "iteration when native camera tooling is not required.",
    )
    add_heading_para(doc, "4.4.2 Key screens and behaviours", 3)
    add_table(
        doc,
        ["Feature", "Behaviour"],
        [
            ["Device connection", "Hub IP + API token; health check"],
            ["Home / dashboard", "Today’s doses, adherence-oriented stats"],
            ["Schedules", "AM/PM picker; long-press edit/delete; block edit after TAKEN today where enforced"],
            ["Reminders", "Poller; “Time Is Up”; nested navigate MainApp → Verify"],
            ["Verify", "Camera preview via hub stream/snapshot; disabled when no actionable dose"],
            ["Monitor", "Missed counts aligned with schedule + grace logic"],
            ["Enrolment UI", "Face (and voice when enabled) against hub enrol endpoints"],
        ],
    )

    add_heading_para(doc, "4.5 Enrolment procedures", 2)
    add_heading_para(doc, "4.5.1 Face enrolment", 3)
    add_para(
        doc,
        "Create or select the patient user; ensure camera preview is live; call face enrol and capture "
        "the configured sample count under good lighting; confirm enrol status; validate with a dry "
        "verify before loading medication.",
    )
    add_heading_para(doc, "4.5.2 Voice enrolment (optional)", 3)
    add_para(
        doc,
        "Enable voice in config and confirm Voice HAT / mic path; fetch challenge text; record and "
        "submit voice enrol; confirm status; include voice in subsequent dispense verify trials.",
    )
    add_para(doc, "[Insert Figure 4.5 — enrolment sequence diagram]", "Caption")

    add_heading_para(doc, "4.6 Mechanical construction and CAD", 2)
    add_heading_para(doc, "4.6.1 Compartment stack", 3)
    add_para(
        doc,
        "Sized Blender models define a compartment envelope on the order of outer diameter ≈ 200 mm "
        "and layer height ≈ 21.67 mm. Six layers correspond to six patients; each cylinder presents "
        "nine equal sectors for doses.",
    )
    add_heading_para(doc, "4.6.2 Drive gear and pinion", 3)
    add_para(
        doc,
        "Layer gear geometry is approximately 72 teeth at 5° tooth spacing on the large ring. A matched "
        "16-tooth stub pinion was designed so pitch radii sum (plus small clearance, e.g. ~0.2 mm) set "
        "centre distance, with tip/root diameters suited to FDM print tolerance.",
    )
    add_para(doc, "[Insert Figure 4.6 — compartment dimensions]", "Caption")
    add_para(doc, "[Insert Figure 4.7 — 72T ring vs 16T pinion mesh]", "Caption")
    add_heading_para(doc, "4.6.3 Assembly notes", 3)
    add_para(
        doc,
        "Print tolerance on tooth thickness and backlash dominates continuous-servo open-loop accuracy; "
        "mechanical calibration and degrees_per_second must be co-tuned. Verify gravity drop path and "
        "IR sight lines before locking the outer shell.",
    )

    add_heading_para(doc, "4.7 Configuration and dry-run practice", 2)
    add_para(
        doc,
        "All tuneables live in config.yaml (face thresholds, grace period, servo continuous calibration, "
        "buzzer timing, GSM port, API token). Unit tests and simulation paths support bring-up without "
        "full mechanics. Commissioning includes changing the default API token and enabling the serial "
        "port while disabling the serial console login.",
    )

    add_heading_para(doc, "4.8 Chapter summary", 2)
    add_para(
        doc,
        "PillSafe was implemented as a Pi 5 hub with continuous MG996R slot indexing, Flask APIs "
        "including camera and biometrics, SIM800C UART alerts, and a React Native client with schedule "
        "and reminder UX. Mechanical CAD progressed to sized layers with an explicit gear–pinion pair. "
        "The next chapter evaluates these subsystems through structured tests.",
    )

    # --- Chapter 5 ---
    add_heading_para(doc, "Chapter 5 - Testing and Results", 1)
    add_heading_para(doc, "5.1 Introduction", 2)
    add_para(
        doc,
        "This chapter reports functional and performance tests of PillSafe against the project "
        "objectives: biometric access control, timed dispensing, reminder behaviour, missed-dose "
        "handling, and mobile schedule/monitoring features. Tests were run on the Raspberry Pi 5 hub "
        "with the React Native client unless noted. Where physical GSM or IR hardware was unavailable, "
        "those paths were exercised in simulation or with the dose-due test endpoint and logged as such. "
        "Quantitative cells below are marked Pending measurement pending completion of the formal "
        "test campaign.",
    )

    add_heading_para(doc, "5.2 Test environment", 2)
    add_table(
        doc,
        ["Item", "Configuration"],
        [
            ["Hub", "Raspberry Pi 5, Raspberry Pi OS 64-bit"],
            ["Config", "hardware/config.yaml as deployed for the trial"],
            ["Face model", "MobileFaceNet TFLite + Haar cascade"],
            ["Servo mode", "Continuous; angle_per_slot 40°; calibrated degrees_per_second"],
            ["App", "React Native build / Expo Go as used in demo"],
            ["Network", "Pi hotspot or LAN; API token set"],
            ["GSM", PENDING + " (SIM800C live or AT simulation)"],
            ["IR", PENDING + " (required true/false per table)"],
        ],
    )

    add_heading_para(doc, "5.3 Facial verification", 2)
    add_para(
        doc,
        "Procedure: Enrol one authorised user under controlled lighting. Run genuine-user and "
        "impostor/unenrolled verify attempts. Record accept/reject and mean confidence where available.",
    )
    add_para(doc, "Table 5.1 — Face verification", "Caption")
    add_table(
        doc,
        ["Trial class", "Attempts (N)", "Accepted", "Rejected", "Rate", "Notes"],
        [
            ["Genuine user", PENDING, PENDING, PENDING, "TAR = " + PENDING, PENDING],
            ["Impostor / unknown", PENDING, PENDING, PENDING, "FAR = " + PENDING, PENDING],
            ["Poor lighting / angle", PENDING, PENDING, PENDING, PENDING, "Optional stress"],
        ],
    )
    add_para(
        doc,
        "Discussion: Comment on confidence_threshold / distance_threshold, retry sets "
        "(max_retries, reject_sets_before_sms), and any lockouts observed once measurements are complete.",
    )
    add_para(doc, "[Insert Figure 5.1 — bar chart of accept/reject]", "Caption")

    add_heading_para(doc, "5.4 Voice verification", 2)
    add_para(
        doc,
        "Procedure: Enrol voice for the same user and run genuine and impostor (or wrong phrase) trials "
        "with the Voice HAT. If voice was disabled in the assessed build, mark results as deferred.",
    )
    add_para(doc, "Table 5.2 — Voice verification", "Caption")
    add_table(
        doc,
        ["Trial class", "Attempts (N)", "Accepted", "Rejected", "Rate", "Notes"],
        [
            ["Genuine", PENDING, PENDING, PENDING, PENDING, PENDING],
            ["Impostor / wrong phrase", PENDING, PENDING, PENDING, PENDING, PENDING],
            ["Face+voice combined policy", PENDING, PENDING, PENDING, PENDING, "If both required"],
        ],
    )

    add_heading_para(doc, "5.5 Dispense mechanics", 2)
    add_para(
        doc,
        "Procedure: After successful verify (or dispenser unit test / dry-run), command one-slot advances. "
        "Measure time per step and visual/angular success (≈40°). With IR enabled, record drop/pickup detection.",
    )
    add_para(doc, "Table 5.3 — Slot advance", "Caption")
    add_table(
        doc,
        ["Compartment", "Steps (N)", "Correct ≈40°", "Overshoot / undershoot", "Mean time (s)", "Failures"],
        [
            ["1 (pin 12)", PENDING, PENDING, PENDING, PENDING, PENDING],
            ["2 (pin 13)", PENDING, PENDING, PENDING, PENDING, PENDING],
            ["3 (pin 22)", PENDING, PENDING, PENDING, PENDING, PENDING],
            ["4 (pin 17)", PENDING, PENDING, PENDING, PENDING, PENDING],
            ["5 (pin 26)", PENDING, PENDING, PENDING, PENDING, PENDING],
            ["6 (pin 27)", PENDING, PENDING, PENDING, PENDING, PENDING],
            ["Overall", PENDING, PENDING, PENDING, PENDING, PENDING],
        ],
    )
    add_para(doc, "Table 5.4 — IR confirmation", "Caption")
    add_table(
        doc,
        ["Event", "Trials (N)", "Detected", "Missed", "Detect rate"],
        [
            ["Pill drop (pin 23)", PENDING, PENDING, PENDING, PENDING],
            ["Pickup (pin 24)", PENDING, PENDING, PENDING, PENDING],
        ],
    )
    add_para(
        doc,
        "Discussion: Continuous open-loop drive depends on degrees_per_second and gear backlash; "
        "report calibration method and residual error once measured.",
    )

    add_heading_para(doc, "5.6 Schedule → reminder → app alert latency", 2)
    add_para(
        doc,
        "Procedure: Create a schedule shortly ahead of wall-clock time. Measure hub recognition of "
        "dose-due (buzzer start), app “Time Is Up” appearance, and end-to-end latency. "
        "POST /alerts/dose-due-test may be used for buzzer-only lab checks.",
    )
    add_para(doc, "Table 5.5 — Reminder latency", "Caption")
    add_table(
        doc,
        ["Run", "Scheduled time", "Buzzer start Δt (s)", "App alert Δt (s)", "Verify Now reachable?"],
        [
            ["1", PENDING, PENDING, PENDING, PENDING],
            ["2", PENDING, PENDING, PENDING, PENDING],
            ["3", PENDING, PENDING, PENDING, PENDING],
            ["Mean", "—", PENDING, PENDING, "—"],
        ],
    )

    add_heading_para(doc, "5.7 Missed-dose path", 2)
    add_para(
        doc,
        "Procedure: Allow a due dose to expire past grace_period_minutes without successful verify. "
        "Confirm DB status MISSED, caregiver SMS (or simulated send log), and Monitor / Home missed counts.",
    )
    add_para(doc, "Table 5.6 — Missed-dose handling", "Caption")
    add_table(
        doc,
        ["Run", "Grace (min)", "MISSED in DB?", "SMS sent / logged?", "App Monitor count matches?"],
        [
            ["1", "15", PENDING, PENDING, PENDING],
            ["2", PENDING, PENDING, PENDING, PENDING],
        ],
    )

    add_heading_para(doc, "5.8 Mobile application behaviour", 2)
    add_para(doc, "Table 5.7 — App functional checks", "Caption")
    add_table(
        doc,
        ["Feature", "Pass/Fail", "Notes"],
        [
            ["Connect with token + health", PENDING, PENDING],
            ["Create schedule (AM/PM)", PENDING, PENDING],
            ["Long-press edit schedule", PENDING, PENDING],
            ["Long-press delete schedule", PENDING, PENDING],
            ["Edit blocked after TAKEN today (if enforced)", PENDING, PENDING],
            ["Reminder → navigate to Verify", PENDING, PENDING],
            ["Verify disabled when no actionable dose", PENDING, PENDING],
            ["Monitor missed vs Home consistency", PENDING, PENDING],
            ["Camera preview (stream/snapshot)", PENDING, PENDING],
        ],
    )

    add_heading_para(doc, "5.9 Limitations observed", 2)
    add_para(
        doc,
        "1. Continuous-servo calibration — open-loop timing drifts with supply voltage, load, and print "
        "backlash; without an encoder, absolute slot index can drift over many cycles.\n"
        "2. Background notifications — when the app is fully backgrounded or killed, OS limits may "
        "prevent reminder polls; users should keep the app available during dose windows or rely on "
        "the hub buzzer and SMS.\n"
        "3. Gear print tolerance — FDM tooth geometry affects mesh and step consistency.\n"
        "4. IR optional mode — if required is false, TAKEN can be logged without physical confirmation.\n"
        "5. GSM dependency — SMS needs network, SIM credit, and correct UART level shifting.",
    )

    add_heading_para(doc, "5.10 Chapter summary", 2)
    add_para(
        doc,
        "Quantitative highlights (face TAR/FAR, dispense success rate, mean reminder latency, missed-path "
        "correctness) will be summarised here once Pending measurement cells are replaced with campaign "
        "data. Objectives demonstrated in software and construction are documented in Chapters 3–4; "
        "formal measured results remain to be completed where marked Pending measurement.",
    )

    # --- Chapter 6 ---
    add_heading_para(doc, "Chapter 6 - Conclusion and Future Work", 1)
    add_heading_para(doc, "6.1 Conclusion", 2)
    add_para(
        doc,
        "This project set out to design and implement PillSafe, a biometric, schedule-driven medication "
        "dispenser for home and small-clinic use. The realised system centres on a Raspberry Pi 5 hub "
        "and a React Native mobile application.",
    )
    add_para(
        doc,
        "On the hub, prescribed times are stored in SQLite and supervised by a schedule controller. When "
        "a dose becomes due, the device issues a buzzer reminder and the phone can surface a “Time Is Up” "
        "alert. Medication is released only after successful biometric checks—facial recognition with "
        "MobileFaceNet TFLite, with an optional voice enrolment/verification path. Dispensing uses six "
        "MG996R continuous-rotation servos, each advancing a nine-slot cylinder in 40° calibrated steps, "
        "with gravity delivery and optional IR confirmation. Caregivers can be notified by SMS through a "
        "SIM800C on the Pi UART when doses are missed or verification fails repeatedly.",
    )
    add_para(
        doc,
        "Relative to the early design draft, the implementation clarified several engineering choices: "
        "Pi 5 GPIO compatibility, continuous rather than absolute positional servo control, per-compartment "
        "motors instead of a single SG90 carousel step, UART GSM instead of USB-serial SIM800L as the "
        "primary path, and React Native instead of Flutter. Mechanical CAD progressed to sized compartment "
        "geometry and a matched ~72T / 16T gear–pinion pair.",
    )
    add_para(
        doc,
        "Overall, PillSafe demonstrates a coherent end-to-end pipeline—from schedule and reminder, through "
        "identity-gated dispense, to adherence logging and caregiver alert—suitable as a final-year "
        "prototype and as a foundation for further clinical hardening.",
    )

    add_heading_para(doc, "6.2 Objectives versus achievements", 2)
    add_table(
        doc,
        ["#", "Objective", "Status", "Evidence"],
        [
            ["1", "Multi-compartment dispenser with identity control", "Achieved", "Ch.3–4; 6×9 mechanism; face (± voice)"],
            ["2", "Scheduling and dose reminders", "Achieved", "Schedule controller; buzzer; app poller"],
            ["3", "Caregiver notify on missed / failed access", "Partially achieved", "SIM800C path; DB MISSED; Ch.5 Table 5.6 (" + PENDING + ")"],
            ["4", "Mobile interface for management", "Achieved", "React Native: connect, schedules, verify, monitor"],
            ["5", "Evaluate system performance", "In progress", "Ch.5 tables — " + PENDING],
        ],
    )

    add_heading_para(doc, "6.3 Contributions", 2)
    add_para(
        doc,
        "1. An integrated Pi 5 hub with config-driven face, schedule, dispense, buzzer, and GSM modules.\n"
        "2. A continuous-servo slot index model aligned with a nine-slot / 40° mechanical layout.\n"
        "3. A React Native client with connection token, AM/PM schedules, edit/delete, reminder UX, and verify handshake.\n"
        "4. Optional voice enrolment APIs and camera streaming for phone-assisted biometrics.\n"
        "5. CAD and wiring documentation (BOM, wiring guide, sized gear/pinion) supporting reproducible construction.",
    )

    add_heading_para(doc, "6.4 Limitations", 2)
    add_para(
        doc,
        "Open-loop continuous servos without encoders; reminder reliability when the mobile OS suspends "
        "the app; IR and GSM sometimes optional or simulated during lab bring-up; prototype enclosure / "
        "full six-layer print validation may still be incomplete; not a certified medical device; no "
        "clinical trial was conducted within this FYP scope.",
    )

    add_heading_para(doc, "6.5 Future work", 2)
    add_para(
        doc,
        "1. Push notifications (FCM/APNs) so caregivers and patients are alerted when the app is backgrounded.\n"
        "2. Full six-layer print and soak test — multi-day adherence run with all compartments loaded.\n"
        "3. Closed-loop slot indexing — magnetic encoder or limit switch per cylinder to eliminate drift.\n"
        "4. Optional cloud sync — encrypted backup of schedules and adherence for multi-caregiver access.\n"
        "5. Clinical usability study — elderly users and caregivers; measure adherence lift versus manual organisers.\n"
        "6. Hardening: rotated API tokens, TLS on the local API, tamper detection on the enclosure.",
    )

    add_heading_para(doc, "6.6 Final remarks", 2)
    add_para(
        doc,
        "PillSafe shows that commodity SBC hardware, careful GPIO and power design, and a focused mobile "
        "client can deliver identity-gated dispensing with caregiver visibility. Completing the measured "
        "test campaign in Chapter 5 leaves the dissertation aligned with the working prototype.",
    )


def update_toc_chapter_entries(doc: Document) -> None:
    """Best-effort update of static TOC lines for chapters 4–6."""
    for p in doc.paragraphs:
        t = p.text.strip()
        if t.startswith("CHAPTER 4") or "CHAPTER 4 -" in t.upper():
            set_runs_text(p, "CHAPTER 4 - IMPLEMENTATION AND CONSTRUCTION")
        # Insert is hard in TOC field; append note paragraphs after last chapter toc if found
    # Find last chapter toc and add after chapter 3 summary if chapter 4 missing as implementation
    insert_after = None
    for i, p in enumerate(doc.paragraphs):
        if "3.5" in p.text and "CHAPTER SUMMARY" in p.text.upper():
            insert_after = p
    if insert_after is None:
        return
    # Only add if CHAPTER 4 implementation not already in toc region
    toc_texts = " ".join(p.text for p in doc.paragraphs if p.style and "toc" in p.style.name.lower())
    if "IMPLEMENTATION AND CONSTRUCTION" in toc_texts.upper():
        return

    def insert_after_paragraph(paragraph: Paragraph, text: str, style_name: str) -> Paragraph:
        new_p = deepcopy(paragraph._element)
        # Clear text in clone then we'll set via paragraph wrapper - simpler: use XML insert
        new_el = paragraph._element.makeelement(paragraph._element.tag, paragraph._element.attrib)
        paragraph._element.addnext(new_el)
        new_para = Paragraph(new_el, paragraph._parent)
        try:
            new_para.style = style_name
        except Exception:
            pass
        set_runs_text(new_para, text)
        return new_para

    # Use Normal with toc-like text if toc style available
    style = insert_after.style.name if insert_after.style else "toc 1"
    p4 = insert_after_paragraph(insert_after, "CHAPTER 4 - IMPLEMENTATION AND CONSTRUCTION", style)
    p5 = insert_after_paragraph(p4, "CHAPTER 5 - TESTING AND RESULTS", style)
    insert_after_paragraph(p5, "CHAPTER 6 - CONCLUSION AND FUTURE WORK", style)


def add_mobile_features(doc: Document) -> None:
    for p in doc.paragraphs:
        if p.text.strip() == "Caregiver interaction":
            # Insert additional feature bullets after this paragraph by appending runs in following empties - simpler rewrite list
            break
    # Expand feature list paragraph block by rewriting known short list items via global context
    replacements_local = {
        "User registration": "User registration / device connection with API token",
        "Medication schedule management": "Medication schedule management (AM/PM picker; long-press edit/delete)",
        "Adherence monitoring": "Adherence monitoring (Home and Monitor aligned with grace logic)",
        "Notification viewing": "Notification viewing and dose-due “Time Is Up” reminder alert",
        "Caregiver interaction": "Caregiver visibility via SMS alerts and in-app notifications; Verify Now handshake",
    }
    for p in doc.paragraphs:
        t = p.text.strip()
        if t in replacements_local:
            set_runs_text(p, replacements_local[t])


def extend_software_modules(doc: Document) -> None:
    # After Flask REST API list item, voice already covered in narrative; add if exact list
    for p in doc.paragraphs:
        if p.text.strip() == "Facial recognition module":
            set_runs_text(p, "Facial recognition module (MobileFaceNet TFLite)")
        if p.text.strip() == "Dispensing controller":
            set_runs_text(p, "Dispensing controller (continuous MG996R slot indexing)")
        if p.text.strip() == "GSM communication service":
            set_runs_text(p, "GSM communication service (SIM800C UART)")
        if p.text.strip() == "Mobile application interface":
            set_runs_text(p, "Mobile application interface (React Native / Expo Go)")


def main() -> None:
    doc = Document(DOC_PATH)

    # 1) Declaration: Bight → Bright
    for p in doc.paragraphs:
        if "Bight" in (p.text or ""):
            set_runs_text(p, p.text.replace("Bight", "Bright"))

    # 2) Abstract
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == "Abstract":
            # next non-empty normal is abstract body
            for j in range(i + 1, min(i + 5, len(doc.paragraphs))):
                if doc.paragraphs[j].text.strip():
                    set_runs_text(doc.paragraphs[j], ABSTRACT.replace("\n\n", " "))
                    # Prefer multi-paragraph abstract: split into 4 paras if following empties exist
                    parts = ABSTRACT.split("\n\n")
                    set_runs_text(doc.paragraphs[j], parts[0])
                    # fill following empty paragraphs if available
                    empties = []
                    for k in range(j + 1, min(j + 8, len(doc.paragraphs))):
                        if not doc.paragraphs[k].text.strip() and doc.paragraphs[k].style.name == "Normal":
                            empties.append(doc.paragraphs[k])
                        elif doc.paragraphs[k].text.strip():
                            break
                    for idx, part in enumerate(parts[1:]):
                        if idx < len(empties):
                            set_runs_text(empties[idx], part)
                        else:
                            # insert after last used
                            break
                    break
            break

    # 3) Specific Ch.3 narrative rewrites (before broad replace)
    n_spec = apply_specific_rewrites(doc)
    print(f"Specific rewrites applied: {n_spec}")

    # 4) Tables
    update_comparison_table(doc)
    update_hardware_table(doc)

    # 5) Software module / mobile feature tweaks
    extend_software_modules(doc)
    add_mobile_features(doc)

    # 6) Broad leftover replacements (order matters — longer first)
    leftovers = [
        ("Raspberry Pi 4 Model B", "Raspberry Pi 5"),
        ("Raspberry Pi 4B", "Raspberry Pi 5"),
        ("raspberry pi 4B", "Raspberry Pi 5"),
        ("Raspberry pi 4B", "Raspberry Pi 5"),
        ("Pi 4B", "Pi 5"),
        ("Flutter development platform using Dart programming language", "React Native framework"),
        ("built with Flutter", "built with React Native"),
        ("built on the Flutter", "built on React Native"),
        ("Flutter", "React Native"),
        ("Dart programming language", "TypeScript/JavaScript"),
        ("SG90 Servo Motor", "MG996R continuous-rotation servo"),
        ("SG90 servo motor", "MG996R continuous-rotation servo"),
        ("SG90 servo", "MG996R continuous-rotation servo"),
        ("SG90", "MG996R"),
        ("SIM800L GSM Module", "SIM800C GSM Module"),
        ("SIM800L GSM module", "SIM800C GSM module"),
        ("SIM800L GPRS GSM Module", "SIM800C GSM Module"),
        ("SIM800L", "SIM800C"),
        ("USB-to-serial communication", "UART (/dev/serial0) communication"),
        ("Serial USB connection interface", "UART (/dev/serial0) interface"),
        ("GPIO18 pin", "servo PWM pins (BCM 12, 13, 22, 17, 26, 27)"),
        ("GPIO18", "BCM 12/13/22/17/26/27"),
        ("steps of 60 degrees", "steps of 40 degrees"),
        ("60 degrees", "40 degrees"),
        ("PCA9685 16-Channel PWM Servo Driver Board", "direct GPIO PWM (no PCA9685 required)"),
        ("PCA9685 16-Channel PWM Servo Driver", "direct GPIO PWM (no PCA9685 required)"),
        ("through the PCA9685", "via direct GPIO PWM rather than a PCA9685"),
        ("PCA9685", "direct GPIO PWM (PCA9685 not required)"),
        ("facial verification mechanism", "facial and optional voice verification mechanism"),
        ("Facial Recognition", "Facial + optional voice"),
    ]
    n_glob = global_replace_in_paragraphs(doc, leftovers)
    print(f"Paragraphs/cells touched by global replace: {n_glob}")

    # 7) Replace stub Chapter 4 and append Ch.4–6
    clear_from_chapter4(doc)
    append_chapters_4_to_6(doc)

    # 8) TOC chapter titles (best effort)
    update_toc_chapter_entries(doc)

    # 9) Heading title for 3.4.4 if still FaceNet-only
    for p in doc.paragraphs:
        if "3.4.4" in p.text and "FaceNet" in p.text:
            set_runs_text(p, "3.4.4 MobileFaceNet (TFLite) Verification Software")
        if p.text.strip().startswith("Chapter 3") and "Methodology" in p.text:
            pass

    doc.save(DOC_PATH)
    print("Saved:", DOC_PATH)

    # Verify
    doc2 = Document(DOC_PATH)
    texts = "\n".join(p.text for p in doc2.paragraphs)
    checks = {
        "Pi 5": "Raspberry Pi 5" in texts,
        "React Native": "React Native" in texts,
        "MG996R": "MG996R" in texts,
        "SIM800C": "SIM800C" in texts,
        "Bright": "Bright" in texts and "Bight" not in texts,
        "Ch4 impl": "Chapter 4 - Implementation" in texts,
        "Ch5": "Chapter 5 - Testing" in texts,
        "Ch6": "Chapter 6 - Conclusion" in texts,
        "Pending": PENDING in texts,
        "no Flutter": "Flutter" not in texts,
        "no SG90": "SG90" not in texts,
        "no Pi 4B": "Pi 4B" not in texts and "Raspberry Pi 4B" not in texts,
    }
    for k, v in checks.items():
        print(f"CHECK {k}: {'OK' if v else 'FAIL'}")


if __name__ == "__main__":
    main()
